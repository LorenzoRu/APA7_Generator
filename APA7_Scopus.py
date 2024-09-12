import pandas as pd
import docx as dc
import tkinter as tk
from tkinter import filedialog, messagebox

# Formatage des auteurs


def format_authors(authors_str):
    authors_list = authors_str.split(',')
    if len(authors_list) == 1:
        return authors_list[0]
    elif len(authors_list) <= 20:
        return ', '.join(authors_list[:-1]) + ' & ' + f'{authors_list[-1]}'
    else:
        sorted_authors = ", ".join(authors_list[:19])
        last_author = authors_list[-1]
        return f"{sorted_authors}, ..., {last_author}"


# Génération du fichier "bibliographie APA7"
def generate_APAlib():
    csv_file = filedialog.askopenfilename(
        title='Ouvrir le fichier CSV', filetypes=[('Fichier CSV', '*.csv')])

    if not csv_file:
        messagebox.showwarning('Erreur', 'Aucun fichier sélectionné')
        return

    # Lecture du fichier CSV
    try:
        df = pd.read_csv(csv_file)
    except Exception as e:
        messagebox.showerror('Erreur', f'Erreur de lecture du fichier CSV: {e}')
        return

    # Création du fichier docx
    doc = dc.Document()
    doc.add_heading('Bibliographie', level=1)

    for i, row in df.iterrows():
        authors = format_authors(row['Authors'])
        title = row['Title']
        year = row['Year']
        url = row['Link']

        # Formatage de la référence
        reference = f"{authors} ({year}). {title}. {url}"

        # Ajout de la référence dans le fichier docx
        doc.add_paragraph(reference)

    # Enregistrement du fichier docx
    doc_file = filedialog.asksaveasfilename(
        title='Enregistrer le fichier', filetypes=[('Fichier Word', '*.docx')])
    if not doc_file:
        messagebox.showwarning('Erreur', 'Aucune destination sélectionné')
        return
    
    try:
        doc.save(doc_file)
        messagebox.showinfo('Succès', f"Le fichier a été enregistré avec succès sous le nom '{doc_file}'")
    except Exception as e:
        messagebox.showerror('Erreur', f"Erreur d'enregistrement du fichier: {e}")


# Interface graphique
def main():
    window = tk.Tk()
    window.title('Générateur de bibliographie format APA7')
    window.geometry('400x200')

    # Bouton pour générer la bibliographie
    generate_button = tk.Button(window, text='Générer la bibliographie', command=generate_APAlib)
    generate_button.pack(pady=50)
    window.mainloop()

main()